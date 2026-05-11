"""
Document Processor - Handles text extraction from various file formats
Supports: PDF, DOCX, CSV, TXT, XLSX
"""
import os
from pathlib import Path
from typing import List, Dict, Optional
import config
from core.logger import get_logger
from core.exceptions import UnsupportedFileTypeError, FileExtractionError

logger = get_logger(__name__)


class DocumentProcessor:
    """Process and extract text from documents"""
    
    def __init__(self):
        self.chunk_size = config.CHUNK_SIZE
        self.chunk_overlap = config.CHUNK_OVERLAP
    
    def process_file(self, file_path: str) -> Dict:
        """Process a file and extract its text content"""
        path = Path(file_path)
        file_type = path.suffix.lower().lstrip('.')
        
        logger.info(f"Processing file: {path.name} (type: {file_type})")
        
        if file_type == 'pdf':
            text = self._extract_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            text = self._extract_docx(file_path)
        elif file_type in ['csv']:
            text = self._extract_csv(file_path)
        elif file_type in ['xlsx', 'xls']:
            text = self._extract_excel(file_path)
        elif file_type == 'txt':
            text = self._extract_txt(file_path)
        else:
            logger.error(f"Unsupported file type: {file_type}")
            raise UnsupportedFileTypeError(file_type, list(config.ALLOWED_EXTENSIONS))
        
        # Create chunks
        chunks = self._create_chunks(text, path.name)
        
        # ── NEW: Extract Knowledge Graph ──
        try:
            from services.agent.rag.graph_extractor import graph_extractor
            logger.info(f"Automatically extracting Knowledge Graph for {path.name}...")
            graph_data = graph_extractor.extract_graph(text)
        except Exception as e:
            logger.error(f"Automatic graph extraction failed: {e}")
            graph_data = {"nodes": [], "edges": []}
        
        logger.info(f"Successfully processed {path.name}: {len(chunks)} chunks created")
        
        return {
            'filename': path.name,
            'file_type': file_type,
            'text': text,
            'chunks': chunks,
            'chunk_count': len(chunks),
            'graph_data': graph_data
        }
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple extractors for robustness"""
        # Method 1: pdfminer.six (superior for Arabic/Complex layouts)
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(file_path)
            if text and len(text.strip()) > 50:
                logger.debug(f"Successfully extracted {len(text)} chars using pdfminer")
                return text.strip()
        except Exception as e:
            logger.debug(f"pdfminer failed, trying fallback: {e}")

        # Method 2: PyPDF2 (Faster, fallback)
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
            
            result = "\n\n".join(text_parts).strip()
            if result:
                return result
        except Exception as e:
            logger.error(f"PyPDF2 failed for {file_path}: {e}")

        # If all failed
        logger.error(f"Failed to extract text from PDF {file_path}")
        return ""
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Failed to extract DOCX {file_path}: {e}")
            raise FileExtractionError(file_path, 'docx', str(e))
    
    def _extract_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Create a readable text representation
            text_parts = []
            text_parts.append(f"CSV File with {len(df)} rows and {len(df.columns)} columns")
            column_names = [str(c) for c in df.columns]
            text_parts.append(f"Columns: {', '.join(column_names)}")
            text_parts.append("\nData:")
            
            # Convert each row to readable text
            for i, (idx, row) in enumerate(df.iterrows()):
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(f"Row {i + 1}: {row_text}")
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract CSV {file_path}: {e}")
            raise FileExtractionError(file_path, 'csv', str(e))
    
    def _extract_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                # Read with a row limit to prevent OOM on massive sheets
                df = pd.read_excel(excel_file, sheet_name=sheet_name, nrows=5000)
                text_parts.append(f"\n[Sheet: {sheet_name}]")
                text_parts.append(f"Rows: {len(df)}, Columns: {len(df.columns)}")
                column_names = [str(c) for c in df.columns]
                text_parts.append(f"Columns: {', '.join(column_names)}")
                
                for i, (idx, row) in enumerate(df.iterrows()):
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                    text_parts.append(f"Row {i + 1}: {row_text}")
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract Excel {file_path}: {e}")
            raise FileExtractionError(file_path, 'xlsx', str(e))
    
    def _extract_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Failed to extract TXT {file_path}: {e}")
            raise FileExtractionError(file_path, 'txt', str(e))
    
    def _create_chunks(self, text: str, filename: str) -> List[Dict]:
        """Split text into semantic chunks based on structure"""
        if not text.strip():
            return []
            
        import re
        
        # Section header patterns: Markdown (#), Numbered (1. Title), ALL CAPS
        SECTION_HEADER_RE = re.compile(
            r"^(?:"
            r"#{1,6}\s+.+|"             # Markdown headers
            r"\d+[\.\)]\s+[A-Z\u0600-\u06FF].+|"  # Numbered items
            r"[A-Z\u0600-\u06FF][A-Z\u0600-\u06FF\s]{5,}$|"  # ALL CAPS line
            r"[^\n]{2,50}:\s*$"         # Short line ending with colon
            r")",
            re.MULTILINE
        )
        
        # Find all headers
        headers = list(SECTION_HEADER_RE.finditer(text))
        
        chunks = []
        chunk_index = 0
        
        # Helper to add a chunk
        def add_chunk(content, index, metadata_suffix=""):
            if len(content.strip()) > 50: # Ignore tiny chunks
                chunks.append({
                    'index': index,
                    'content': content.strip(),
                    'metadata': f"Source: {filename}, Chunk: {index + 1} {metadata_suffix}".strip()
                })
                return index + 1
            return index

        # If no headers found, fallback to paragraph-based chunking
        if not headers:
            # Fallback logic (original implementation)
            paragraphs = text.split('\n\n')
            current_chunk = ""
            
            for para in paragraphs:
                para = para.strip()
                if not para: continue
                
                if len(current_chunk) + len(para) > self.chunk_size:
                    if current_chunk:
                        chunk_index = add_chunk(current_chunk, chunk_index)
                        # Keep overlap by taking the last X characters safely
                        overlap_chars = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else ""
                        # Find the first space to avoid cutting words in half
                        if " " in overlap_chars:
                            overlap_chars = overlap_chars.split(" ", 1)[-1]
                        
                        current_chunk = overlap_chars + '\n\n' + para
                    else:
                        current_chunk = para
                else:
                    current_chunk = current_chunk + '\n\n' + para if current_chunk else para
            
            if current_chunk:
                add_chunk(current_chunk, chunk_index)
                
            return chunks

        # Semantic Chunking logic
        current_header = None
        current_content = []
        
        # Initial content before first header
        pre_header_content = text[:headers[0].start()].strip()
        if pre_header_content:
            chunk_index = add_chunk(pre_header_content, chunk_index, "(Intro)")
            
        for i, match in enumerate(headers):
            header_text = match.group().strip()
            start = match.end()
            end = headers[i+1].start() if i + 1 < len(headers) else len(text)
            section_content = text[start:end].strip()
            
            # If section content is small, merge or just add
            full_section = f"[{header_text}]\n{section_content}"
            
            if len(full_section) > self.chunk_size * 2: # Too big, split it
                # Split large section by paragraphs
                paras = section_content.split('\n\n')
                sub_chunk = f"[{header_text}]\n"
                
                for para in paras:
                    if len(sub_chunk) + len(para) > self.chunk_size:
                        chunk_index = add_chunk(sub_chunk, chunk_index, f"({header_text})")
                        sub_chunk = f"[{header_text} (cont.)]\n" + para
                    else:
                        sub_chunk += "\n\n" + para
                
                if sub_chunk:
                    chunk_index = add_chunk(sub_chunk, chunk_index, f"({header_text})")
            else:
                # Good size, add as one semantic chunk
                chunk_index = add_chunk(full_section, chunk_index, f"({header_text})")
                
        return chunks
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(config.ALLOWED_EXTENSIONS)


# Singleton instance
document_processor = DocumentProcessor()
